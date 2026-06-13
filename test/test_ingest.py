#!/usr/bin/env python3
"""文档解析摄取层单测（设计 v3 B批1）。

动态生成测试文件（PDF/DOCX/TXT/MD），不依赖外部 fixture。
覆盖：各类型确定性解析、安全校验、摘要预算、多文件合并、扫描件 needs_vision、错误容错。
纯逻辑 + 真实文件解析（pymupdf/python-docx），任何环境可跑（含 CI）。
"""
from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest

_bs = Path(__file__).resolve().parent / "swarm_bootstrap.py"
_spec = importlib.util.spec_from_file_location("swarm_bootstrap", _bs)
_mod = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_mod)

from swarm.brain import ingest


# ── 测试文件生成 helper ────────────────────────────────────

def _make_txt(tmp_path: Path, name: str, content: str) -> Path:
    p = tmp_path / name
    p.write_text(content, encoding="utf-8")
    return p


def _make_pdf(tmp_path: Path, name: str, text: str) -> Path:
    import fitz
    p = tmp_path / name
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(str(p))
    doc.close()
    return p


def _make_empty_pdf(tmp_path: Path, name: str) -> Path:
    """无文本层的 PDF（模拟扫描件）。"""
    import fitz
    p = tmp_path / name
    doc = fitz.open()
    doc.new_page()  # 空白页，无文本
    doc.save(str(p))
    doc.close()
    return p


def _make_docx(tmp_path: Path, name: str, paragraphs: list[str], table: list[list[str]] | None = None) -> Path:
    import docx
    p = tmp_path / name
    d = docx.Document()
    for para in paragraphs:
        d.add_paragraph(para)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, val in enumerate(row):
                t.rows[i].cells[j].text = val
    d.save(str(p))
    return p


# ── 安全校验 ───────────────────────────────────────────────

def test_validate_rejects_unknown_ext(tmp_path):
    p = _make_txt(tmp_path, "evil.exe", "x")
    err = ingest.validate_file(p)
    assert err and "不支持的格式" in err
    print("  ✅ 安全: 拒绝非白名单扩展名 .exe")


def test_validate_rejects_oversize(tmp_path):
    p = _make_txt(tmp_path, "big.txt", "x" * 2000)
    err = ingest.validate_file(p, max_bytes=1000)
    assert err and "过大" in err
    print("  ✅ 安全: 拒绝超大文件")


def test_validate_rejects_empty(tmp_path):
    p = _make_txt(tmp_path, "empty.txt", "")
    err = ingest.validate_file(p)
    assert err and "空文件" in err
    print("  ✅ 安全: 拒绝空文件")


def test_validate_rejects_missing(tmp_path):
    err = ingest.validate_file(tmp_path / "nope.txt")
    assert err and "不存在" in err
    print("  ✅ 安全: 拒绝不存在文件")


def test_validate_passes_valid(tmp_path):
    p = _make_txt(tmp_path, "ok.md", "# Hello")
    assert ingest.validate_file(p) is None
    print("  ✅ 安全: 合法文件通过")


# ── 各类型解析 ─────────────────────────────────────────────

def test_parse_txt(tmp_path):
    p = _make_txt(tmp_path, "req.txt", "需要做一个登录功能")
    doc = ingest.parse_file(p)
    assert doc.kind == "text"
    assert "登录功能" in doc.text
    assert doc.source == "text"
    assert not doc.error
    print("  ✅ 解析: txt 直读")


def test_parse_md(tmp_path):
    p = _make_txt(tmp_path, "req.md", "# 需求\n- 用户注册\n- 密码重置")
    doc = ingest.parse_file(p)
    assert doc.kind == "text"
    assert "用户注册" in doc.text
    print("  ✅ 解析: md 直读")


def test_parse_pdf_with_text(tmp_path):
    # PDF 用英文文本：默认字体不支持中文（会渲染成豆腐块），测解析逻辑英文等价。
    p = _make_pdf(tmp_path, "spec.pdf", "Requirement: build a shopping cart")
    doc = ingest.parse_file(p)
    assert doc.kind == "pdf"
    assert "shopping cart" in doc.text
    assert doc.page_count == 1
    assert not doc.needs_vision
    print("  ✅ 解析: pdf 文本层抽取")


def test_parse_pdf_scanned_needs_vision(tmp_path):
    p = _make_empty_pdf(tmp_path, "scan.pdf")
    doc = ingest.parse_file(p)
    assert doc.kind == "pdf"
    assert doc.needs_vision is True
    assert doc.source == "none"
    print("  ✅ 解析: 无文本层 pdf → needs_vision (留给B批2)")


def test_parse_docx(tmp_path):
    p = _make_docx(tmp_path, "req.docx",
                   ["项目目标：电商平台", "核心需求如下"],
                   table=[["模块", "优先级"], ["支付", "高"]])
    doc = ingest.parse_file(p)
    assert doc.kind == "docx"
    assert "电商平台" in doc.text
    assert "支付" in doc.text  # 表格内容也抽出
    print("  ✅ 解析: docx 段落+表格抽取")


def test_parse_image_needs_vision(tmp_path):
    # 造一个最小 png
    p = tmp_path / "ui.png"
    p.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
        "890000000a49444154789c6360000002000100" + "0" * 20
    ))
    doc = ingest.parse_file(p)
    assert doc.kind == "image"
    assert doc.needs_vision is True
    print("  ✅ 解析: 图片 → needs_vision (留给B批2)")


def test_parse_error_is_captured(tmp_path):
    # 故意造一个损坏的 pdf（扩展名对但内容不是 pdf）
    p = tmp_path / "broken.pdf"
    p.write_bytes(b"this is not a real pdf content at all")
    doc = ingest.parse_file(p)
    assert doc.error, "损坏文件应记入 error 而非抛出"
    print("  ✅ 解析: 损坏文件 error 被捕获不抛出")


# ── 摘要 / 预算 ────────────────────────────────────────────

def test_summarize_within_budget():
    text = "短文本"
    out, truncated = ingest.summarize_to_budget(text, max_tokens=1000)
    assert out == text and truncated is False
    print("  ✅ 摘要: 预算内不截断")


def test_summarize_truncates():
    text = "A" * 10000
    out, truncated = ingest.summarize_to_budget(text, max_tokens=100)  # 100*4=400 chars
    assert truncated is True
    assert len(out) < len(text)
    assert "已省略" in out
    # 头尾都保留
    assert out.startswith("A")
    assert out.endswith("A")
    print("  ✅ 摘要: 超预算头尾保留+中间省略")


def test_summarize_zero_budget_no_limit():
    text = "A" * 10000
    out, truncated = ingest.summarize_to_budget(text, max_tokens=0)
    assert out == text and truncated is False
    print("  ✅ 摘要: 预算=0 不限制")


# ── 多文件摄取编排 ─────────────────────────────────────────

def test_ingest_multiple_files(tmp_path):
    txt = _make_txt(tmp_path, "a.txt", "文字需求A")
    pdf = _make_pdf(tmp_path, "b.pdf", "PDF requirement B")
    result = ingest.ingest_files([txt, pdf], base_text="总体目标：做个系统")
    assert result.ok
    assert "总体目标" in result.draft_text
    assert "文字需求A" in result.draft_text
    assert "PDF requirement B" in result.draft_text
    assert len(result.documents) == 2
    assert not result.errors
    print("  ✅ 摄取: 多文件+任务描述合并为草稿")


def test_ingest_collects_vision_files(tmp_path):
    txt = _make_txt(tmp_path, "a.txt", "文字需求")
    img = tmp_path / "screen.png"
    img.write_bytes(bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
        "890000000a49444154789c6360000002000100" + "0" * 20
    ))
    result = ingest.ingest_files([txt, img])
    assert "screen.png" in result.needs_vision_files
    assert "待多模态理解" in result.draft_text
    print("  ✅ 摄取: 图片记入 needs_vision_files (留给B批2)")


def test_ingest_continues_on_bad_file(tmp_path):
    good = _make_txt(tmp_path, "good.txt", "好文件")
    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"not a pdf")
    result = ingest.ingest_files([good, bad])
    assert "好文件" in result.draft_text   # 好文件正常
    assert any("bad.pdf" in e for e in result.errors)  # 坏文件记错误但不中断
    print("  ✅ 摄取: 单文件失败不拖垮整体")


def test_ingest_applies_budget(tmp_path):
    big = _make_txt(tmp_path, "big.txt", "需求" * 5000)  # ~10000 字符
    result = ingest.ingest_files([big], context_budget_tokens=100)  # 400 chars
    assert len(result.draft_text) < 10000
    assert result.documents[0].truncated is True
    print("  ✅ 摄取: 超预算时按 context_budget 截断")


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-v", "-s"]))
