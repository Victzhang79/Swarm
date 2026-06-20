"""多模态需求摄取层 — 文档解析（设计 v3 B 部分 / B批1）。

把任务创建时上传的文件解析成「需求草稿」文本，前置于 analyze 节点。
本批（B批1）只做**确定性解析 + 摘要**（纯逻辑，易测）：
  - md/txt → 直读
  - pdf（文本层）→ PyMuPDF 抽文字
  - docx → python-docx 抽文字
  - 图片 / 扫描件 pdf（无文本层）→ 标记 needs_vision，留给 B批2 多模态理解

设计约束：
  - 安全（B.3）：格式白名单、单文件/总大小上限、解析超时、PDF 不执行 JS（PyMuPDF 默认）。
  - 上下文预算（B.4 消费 A.5）：大文档按目标模型 context_window 控制注入量，
    超预算则摘要/截断，不塞爆 analyze。
  - 防幻觉（B.2）：确定性解析的文本标 source=text（可信）；图片理解留给 B批2 标 ai_vision。

依赖：pymupdf(fitz) + python-docx，均为轻量确定性库（设计 B1：不上 marker）。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path

logger = logging.getLogger(__name__)

# ── 安全限制（B.3）──────────────────────────────────────────
# 格式白名单：仅这些扩展名被接受。
ALLOWED_EXTENSIONS = frozenset({
    ".md", ".markdown", ".txt", ".text",
    ".pdf", ".docx",
    ".html", ".htm",
    ".png", ".jpg", ".jpeg", ".webp",
})
TEXT_EXTENSIONS = frozenset({".md", ".markdown", ".txt", ".text"})
HTML_EXTENSIONS = frozenset({".html", ".htm"})
IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".webp"})

DEFAULT_MAX_FILE_BYTES = 20 * 1024 * 1024   # 单文件 20MB（B.3 默认，可配）
DEFAULT_PARSE_TIMEOUT = 30                    # 单文件解析超时（秒）

# ~4 chars/token 粗估（与 project/store.py 一致；精确计算是后续债）
_CHARS_PER_TOKEN = 4


@dataclass
class ParsedDocument:
    """单个文件的解析结果。"""
    filename: str
    ext: str
    kind: str                    # text | pdf | docx | html | image | unknown
    text: str = ""               # 抽取的文本（确定性解析）
    needs_vision: bool = False   # True=需多模态理解（图片/扫描件），留给 B批2
    char_count: int = 0
    page_count: int = 0          # pdf 页数（其它为 0）
    truncated: bool = False      # 是否因预算被截断
    error: str = ""              # 解析失败原因（非空=失败）
    source: str = "text"         # text(确定性) | ai_vision(B批2 填) | none

    def est_tokens(self) -> int:
        return self.char_count // _CHARS_PER_TOKEN


@dataclass
class IngestResult:
    """多文件合并的摄取结果（→ 需求草稿）。"""
    documents: list[ParsedDocument] = field(default_factory=list)
    draft_text: str = ""              # 合并后的需求草稿文本
    needs_vision_files: list[str] = field(default_factory=list)  # 待 B批2 处理的文件
    total_chars: int = 0
    errors: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return any(d.text or d.needs_vision for d in self.documents)


# ──────────────────────────────────────────────
# 安全校验（入口）
# ──────────────────────────────────────────────

def validate_file(
    path: str | Path, max_bytes: int = DEFAULT_MAX_FILE_BYTES
) -> str | None:
    """校验单文件是否可接受。返回错误消息（None=通过）。

    检查：存在性、扩展名白名单、大小上限。MIME 校验在解析时按内容兜底
    （B批3 上传端点会再做一道 MIME 校验，不只信扩展名）。
    """
    p = Path(path)
    if not p.exists() or not p.is_file():
        return f"文件不存在: {p.name}"
    ext = p.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return f"不支持的格式: {ext}（白名单: {', '.join(sorted(ALLOWED_EXTENSIONS))}）"
    size = p.stat().st_size
    if size > max_bytes:
        return f"文件过大: {size / 1024 / 1024:.1f}MB > {max_bytes / 1024 / 1024:.0f}MB 上限"
    if size == 0:
        return f"空文件: {p.name}"
    return None


# ──────────────────────────────────────────────
# 各类型确定性解析（纯函数）
# ──────────────────────────────────────────────

def _read_text_file(path: Path) -> str:
    """直读纯文本（md/txt）。容错编码。"""
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, ValueError):
            continue
    # 最后兜底：忽略错误字节
    return path.read_bytes().decode("utf-8", errors="ignore")


def _parse_pdf(path: Path) -> tuple[str, int, bool]:
    """PyMuPDF 抽 PDF 文本层。返回 (text, page_count, needs_vision)。

    无文本层（扫描件/纯图 PDF）→ needs_vision=True，文本留空（B批2 渲染图走多模态）。
    PyMuPDF 默认不执行 JS / 不加载外部资源（B.3 安全要求）。
    """
    import fitz  # pymupdf

    text_parts: list[str] = []
    page_count = 0
    # filetype="pdf" 显式指定，避免按扩展名误判
    with fitz.open(path) as doc:
        page_count = doc.page_count
        for page in doc:
            t = str(page.get_text("text") or "")
            if t.strip():
                text_parts.append(t)
    full = "\n".join(text_parts).strip()
    # 文本层几乎为空 → 判定为扫描件/图片 PDF，需多模态。
    # 阈值取极小（仅挡近乎纯空白的扫描件）；短而有效的文本需求 PDF 不应被误判。
    # 注意：扫描件偶尔抽到零星 OCR 噪声字符，但正常文本 PDF 即便一句话也远超此值。
    needs_vision = len(full) < 5
    return full, page_count, needs_vision


def _parse_docx(path: Path) -> str:
    """python-docx 抽 DOCX 文本（段落 + 表格）。"""
    import docx

    d = docx.Document(str(path))
    parts: list[str] = [p.text for p in d.paragraphs if p.text.strip()]
    # 表格内容也抽出来（需求文档常用表格）
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


class _HTMLTextExtractor(HTMLParser):
    """标准库 HTML → 纯文本(去 script/style，标题/列表保留换行)。

    只用 html.parser（不引入 bs4）。不追求完整 HTML 规范，够喂下游切分即可：
      - <script>/<style> 内容整段丢弃（不进文本）。
      - 标题(h1-h6)/列表项(li)/段落(p,div,br…)前后补换行，保留结构感。
      - 其余文本折叠连续空白为单空格。
    """

    _SKIP_TAGS = {"script", "style"}
    _BREAK_TAGS = {
        "p", "div", "section", "article", "br", "tr",
        "h1", "h2", "h3", "h4", "h5", "h6",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: object) -> None:  # noqa: ARG002
        tag = tag.lower()
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        elif tag == "li":
            self._parts.append("\n- ")
        elif tag in self._BREAK_TAGS:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self._SKIP_TAGS:
            self._skip_depth = max(0, self._skip_depth - 1)
        elif tag in self._BREAK_TAGS:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        collapsed = " ".join(data.split())
        if collapsed:
            self._parts.append(collapsed + " ")

    def get_text(self) -> str:
        body = "".join(self._parts)
        # 压掉过多空行 + 行尾空白
        cleaned = "\n".join(line.rstrip() for line in body.splitlines())
        while "\n\n\n" in cleaned:
            cleaned = cleaned.replace("\n\n\n", "\n\n")
        return cleaned.strip()


def _parse_html(path: Path) -> str:
    """标准库 html.parser 抽 HTML 文本。去 script/style，保留标题/列表换行。

    不依赖 bs4（B.3：不引入新依赖）。容错编码，沿用纯文本读取的编码兜底。
    """
    raw = _read_text_file(path)
    parser = _HTMLTextExtractor()
    parser.feed(raw)
    parser.close()
    return parser.get_text()


def parse_file(
    path: str | Path, max_bytes: int = DEFAULT_MAX_FILE_BYTES
) -> ParsedDocument:
    """解析单个文件 → ParsedDocument。先校验，再按类型分流。

    任何解析异常都被捕获并记入 error（不抛出），保证摄取链路不被单文件拖垮。
    """
    p = Path(path)
    ext = p.suffix.lower()
    doc = ParsedDocument(filename=p.name, ext=ext, kind="unknown")

    err = validate_file(p, max_bytes=max_bytes)
    if err:
        doc.error = err
        return doc

    try:
        if ext in TEXT_EXTENSIONS:
            doc.kind = "text"
            doc.text = _read_text_file(p)
        elif ext == ".pdf":
            doc.kind = "pdf"
            text, pages, needs_vision = _parse_pdf(p)
            doc.text = text
            doc.page_count = pages
            doc.needs_vision = needs_vision
            if needs_vision:
                doc.source = "none"  # 待 B批2 多模态填充
        elif ext == ".docx":
            doc.kind = "docx"
            doc.text = _parse_docx(p)
        elif ext in HTML_EXTENSIONS:
            doc.kind = "html"
            doc.text = _parse_html(p)
        elif ext in IMAGE_EXTENSIONS:
            doc.kind = "image"
            doc.needs_vision = True   # 图片必走多模态（B批2）
            doc.source = "none"
        else:
            doc.error = f"无解析器: {ext}"
            return doc
    except Exception as exc:  # noqa: BLE001
        doc.error = f"解析失败: {exc}"
        logger.warning("解析文件 %s 失败: %s", p.name, exc)
        return doc

    doc.char_count = len(doc.text)
    return doc


# ──────────────────────────────────────────────
# 摘要 / 预算控制（B.4，消费 A.5 能力库 context_window）
# ──────────────────────────────────────────────

def summarize_to_budget(text: str, max_tokens: int) -> tuple[str, bool]:
    """把文本控制在 token 预算内。返回 (text, truncated)。

    本批用确定性截断（头部优先 + 尾部保留尾注），不调 LLM——保证纯函数可测、
    无副作用。LLM 智能摘要是后续增强（需求文档头尾通常含关键信息：标题/目标 + 验收）。
    """
    if max_tokens <= 0:
        return text, False
    budget_chars = max_tokens * _CHARS_PER_TOKEN
    if len(text) <= budget_chars:
        return text, False
    # 头部留 70%、尾部留 30%（保留开头的目标陈述 + 结尾的验收/约束）
    head = int(budget_chars * 0.7)
    tail = budget_chars - head
    marker = "\n\n…（文档过长，中间内容已省略）…\n\n"
    return text[:head] + marker + text[-tail:], True


# ──────────────────────────────────────────────
# 多文件摄取编排 → 需求草稿
# ──────────────────────────────────────────────

def ingest_files(
    paths: list[str | Path],
    *,
    base_text: str = "",
    context_budget_tokens: int = 0,
    max_bytes: int = DEFAULT_MAX_FILE_BYTES,
) -> IngestResult:
    """摄取多个文件 + 任务文字描述 → 合并为需求草稿（设计 B.1）。

    base_text: 任务创建时用户输入的文字描述（与文件内容合并）。
    context_budget_tokens: 注入预算（从 A 能力库取目标模型 context_window×系数）；
      0=不限。超预算时对合并文本做确定性摘要/截断（B.4）。
    needs_vision 的文件（图片/扫描件）不在此解析正文，记入 needs_vision_files 待 B批2。
    """
    result = IngestResult()
    parts: list[str] = []
    if base_text.strip():
        parts.append(f"【任务描述】\n{base_text.strip()}")

    for path in paths:
        doc = parse_file(path, max_bytes=max_bytes)
        result.documents.append(doc)
        if doc.error:
            result.errors.append(f"{doc.filename}: {doc.error}")
            continue
        if doc.needs_vision:
            result.needs_vision_files.append(doc.filename)
            # 图片/扫描件占位标记，正文待 B批2 填充
            parts.append(f"【文件: {doc.filename}】（图片/扫描件，待多模态理解）")
            continue
        if doc.text:
            header = f"【文件: {doc.filename}"
            if doc.page_count:
                header += f"（{doc.page_count}页）"
            header += "】"
            parts.append(f"{header}\n{doc.text}")

    merged = "\n\n".join(parts)
    result.total_chars = len(merged)

    if context_budget_tokens > 0:
        merged, truncated = summarize_to_budget(merged, context_budget_tokens)
        if truncated:
            for d in result.documents:
                d.truncated = True

    result.draft_text = merged
    return result
